import io
from typing import Literal

from fastapi import APIRouter, Depends, File, HTTPException, Request, UploadFile

from app.rate_limit import limiter
from fastapi.responses import StreamingResponse

from app.models.request import AnonymizeBatchRequest, AnonymizeTextRequest
from app.models.response import (
    AnonymizeTextResponse,
    BatchStatusResponse,
    BatchTaskResponse,
    ProblemDetail,
)
from app.services.anonymize_service import AnonymizeService
from app.tasks.celery_app import process_batch_task
from app.tasks.batch_store import get_batch_status

router = APIRouter(prefix="/anonymize", tags=["anonymize"])


def get_service() -> AnonymizeService:
    from app.main import get_anonymize_service

    return get_anonymize_service()


@router.post(
    "/text",
    response_model=AnonymizeTextResponse,
    responses={400: {"model": ProblemDetail}, 422: {"model": ProblemDetail}},
)
@limiter.limit("100/minute")
async def anonymize_text(
    http_request: Request,
    request: AnonymizeTextRequest,
    service: AnonymizeService = Depends(get_service),
):
    try:
        return await service.anonymize_text(
            request.text,
            strategy=request.strategy,
            entity_types=request.entity_types,
            return_entities=request.return_entities,
            dry_run=request.dry_run,
            operator_id=request.operator_id,
            legal_basis=request.legal_basis,
        )
    except Exception as exc:
        raise HTTPException(
            status_code=400,
            detail={
                "type": "about:blank",
                "title": "Anonymization failed",
                "status": 400,
                "detail": str(exc),
            },
        ) from exc


@router.post("/batch", response_model=BatchTaskResponse)
async def anonymize_batch(request: AnonymizeBatchRequest):
    if len(request.texts) > 1000:
        raise HTTPException(status_code=400, detail="Maximum 1000 texts per batch")
    payload = request.model_dump()
    task = process_batch_task.delay(payload)
    return BatchTaskResponse(
        task_id=task.id,
        status="PENDING",
        message="Batch processing started",
    )


@router.get("/batch/{task_id}", response_model=BatchStatusResponse)
async def get_batch_result(task_id: str):
    status = get_batch_status(task_id)
    if not status:
        raise HTTPException(status_code=404, detail="Task not found")
    return BatchStatusResponse(**status)


@router.post("/document")
async def anonymize_document(
    file: UploadFile = File(...),
    strategy: Literal["REDACT", "PSEUDONYMIZE", "GENERALIZE", "TOKENIZE"] | None = None,
    service: AnonymizeService = Depends(get_service),
):
    content = await file.read()
    filename = file.filename or "document.txt"
    ext = filename.rsplit(".", 1)[-1].lower() if "." in filename else "txt"

    if ext == "txt":
        text = content.decode("utf-8", errors="ignore")
        result = await service.anonymize_text(text, strategy=strategy)
        return StreamingResponse(
            io.BytesIO(result.anonymized_text.encode("utf-8")),
            media_type="text/plain",
            headers={"Content-Disposition": f'attachment; filename="anon_{filename}"'},
        )

    if ext == "docx":
        from docx import Document

        doc = Document(io.BytesIO(content))
        paragraphs = [p.text for p in doc.paragraphs]
        full_text = "\n".join(paragraphs)
        result = await service.anonymize_text(full_text, strategy=strategy)
        out = Document()
        for line in result.anonymized_text.split("\n"):
            out.add_paragraph(line)
        buf = io.BytesIO()
        out.save(buf)
        buf.seek(0)
        return StreamingResponse(
            buf,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="anon_{filename}"'},
        )

    if ext == "pdf":
        from pypdf import PdfReader

        reader = PdfReader(io.BytesIO(content))
        texts = []
        for page in reader.pages:
            texts.append(page.extract_text() or "")
        full_text = "\n".join(texts)
        result = await service.anonymize_text(full_text, strategy=strategy)
        return StreamingResponse(
            io.BytesIO(result.anonymized_text.encode("utf-8")),
            media_type="text/plain",
            headers={"Content-Disposition": f'attachment; filename="anon_{filename}.txt"'},
        )

    raise HTTPException(status_code=415, detail="Supported formats: TXT, DOCX, PDF")
